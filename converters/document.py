import html as html_lib
from pathlib import Path
from html.parser import HTMLParser
from utils.helpers import gerar_nome_saida
from utils.validators import validar_arquivo
from utils.logger import setup_logger

logger = setup_logger(__name__)


class DocumentConverter:
    """Converte entre TXT, PDF, DOCX, HTML, RTF.

    Usa xhtml2pdf (Python puro) em vez de pdfkit/wkhtmltopdf: wkhtmltopdf é
    um binário externo que não empacota sozinho num .exe via PyInstaller
    sem passos extras. xhtml2pdf evita essa dependência.
    """

    def __init__(self):
        self.tipo = 'document'
        self.logger = logger

    def _validar(self, arquivo):
        val = validar_arquivo(arquivo, self.tipo)
        if not val['valido']:
            raise ValueError(val['erro'])
        return Path(arquivo)

    def _html_para_pdf_arquivo(self, html_content, saida):
        from xhtml2pdf import pisa
        with open(saida, 'wb') as f:
            status = pisa.CreatePDF(html_content, dest=f)
        if status.err:
            raise RuntimeError(f"Falha ao gerar PDF ({status.err} erro(s))")

    # ==================== TXT ====================

    def txt_para_pdf(self, arquivo):
        """TXT → PDF"""
        arquivo = self._validar(arquivo)
        conteudo = arquivo.read_text(encoding='utf-8', errors='replace')
        conteudo_escapado = html_lib.escape(conteudo)

        html_content = f"""
        <html><head><meta charset="utf-8"><style>
            @page {{ margin: 2cm; }}
            body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; }}
            pre {{ white-space: pre-wrap; word-wrap: break-word; font-family: Courier, monospace; }}
        </style></head>
        <body><pre>{conteudo_escapado}</pre></body></html>
        """

        saida = gerar_nome_saida(arquivo, "_converted", ".pdf")
        self._html_para_pdf_arquivo(html_content, saida)
        return saida

    def txt_para_docx(self, arquivo):
        """TXT → DOCX"""
        from docx import Document
        arquivo = self._validar(arquivo)
        conteudo = arquivo.read_text(encoding='utf-8', errors='replace')

        doc = Document()
        for linha in conteudo.split('\n'):
            doc.add_paragraph(linha)

        saida = gerar_nome_saida(arquivo, "_converted", ".docx")
        doc.save(str(saida))
        return saida

    # ==================== PDF ====================

    def pdf_para_txt(self, arquivo):
        """PDF → TXT (extrai texto; PDFs escaneados/imagem não têm texto extraível)"""
        from pypdf import PdfReader
        arquivo = self._validar(arquivo)

        reader = PdfReader(str(arquivo))
        texto = "\n".join(page.extract_text() or "" for page in reader.pages)

        if not texto.strip():
            self.logger.warning(f"Nenhum texto extraído de {arquivo} (pode ser PDF escaneado/imagem)")

        saida = gerar_nome_saida(arquivo, "_converted", ".txt")
        saida.write_text(texto, encoding='utf-8')
        return saida

    def pdf_para_docx(self, arquivo):
        """PDF → DOCX (texto puro; não preserva layout/formatação original)"""
        from pypdf import PdfReader
        from docx import Document
        arquivo = self._validar(arquivo)

        reader = PdfReader(str(arquivo))
        doc = Document()
        for i, page in enumerate(reader.pages):
            texto = page.extract_text() or ""
            for linha in texto.split('\n'):
                doc.add_paragraph(linha)
            if i < len(reader.pages) - 1:
                doc.add_page_break()

        saida = gerar_nome_saida(arquivo, "_converted", ".docx")
        doc.save(str(saida))
        return saida

    # ==================== DOCX ====================

    def docx_para_txt(self, arquivo):
        """DOCX → TXT"""
        from docx import Document
        arquivo = self._validar(arquivo)

        doc = Document(str(arquivo))
        texto = "\n".join(p.text for p in doc.paragraphs)

        saida = gerar_nome_saida(arquivo, "_converted", ".txt")
        saida.write_text(texto, encoding='utf-8')
        return saida

    def docx_para_pdf(self, arquivo):
        """DOCX → PDF (parágrafos e formatação básica; não é conversão 1:1 de layout)"""
        from docx import Document
        arquivo = self._validar(arquivo)

        doc = Document(str(arquivo))
        partes_html = []
        for p in doc.paragraphs:
            texto = html_lib.escape(p.text) or "&nbsp;"
            estilo = (p.style.name or "").lower() if p.style else ""
            if "heading" in estilo:
                partes_html.append(f"<h3>{texto}</h3>")
            else:
                partes_html.append(f"<p>{texto}</p>")

        html_content = f"""
        <html><head><meta charset="utf-8"><style>
            @page {{ margin: 2cm; }}
            body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; }}
            p {{ margin: 0 0 8px 0; }}
        </style></head>
        <body>{''.join(partes_html)}</body></html>
        """

        saida = gerar_nome_saida(arquivo, "_converted", ".pdf")
        self._html_para_pdf_arquivo(html_content, saida)
        return saida

    # ==================== HTML ====================

    def html_para_pdf(self, arquivo):
        """HTML → PDF"""
        arquivo = self._validar(arquivo)
        html_content = arquivo.read_text(encoding='utf-8', errors='replace')

        saida = gerar_nome_saida(arquivo, "_converted", ".pdf")
        self._html_para_pdf_arquivo(html_content, saida)
        return saida

    def html_para_txt(self, arquivo):
        """HTML → TXT (extrai apenas texto visível, ignora <script>/<style>)"""
        arquivo = self._validar(arquivo)
        html_content = arquivo.read_text(encoding='utf-8', errors='replace')

        class ExtratorTexto(HTMLParser):
            def __init__(self):
                super().__init__()
                self.partes = []
                self._ignorar = 0

            def handle_starttag(self, tag, attrs):
                if tag in ('script', 'style'):
                    self._ignorar += 1
                elif tag in ('br', 'p', 'div', 'li', 'tr', 'h1', 'h2', 'h3', 'h4'):
                    self.partes.append('\n')

            def handle_endtag(self, tag):
                if tag in ('script', 'style') and self._ignorar > 0:
                    self._ignorar -= 1

            def handle_data(self, data):
                if self._ignorar == 0 and data.strip():
                    self.partes.append(data.strip())

        parser = ExtratorTexto()
        parser.feed(html_content)
        texto = html_lib.unescape('\n'.join(p for p in parser.partes if p))

        # Colapsa linhas em branco excessivas
        linhas = [l for l in texto.split('\n')]
        texto_final = '\n'.join(linhas)

        saida = gerar_nome_saida(arquivo, "_converted", ".txt")
        saida.write_text(texto_final, encoding='utf-8')
        return saida

    # ==================== RTF ====================

    def rtf_para_txt(self, arquivo):
        """RTF → TXT"""
        from striprtf.striprtf import rtf_to_text
        arquivo = self._validar(arquivo)

        conteudo_rtf = arquivo.read_text(encoding='utf-8', errors='replace')
        texto = rtf_to_text(conteudo_rtf)

        saida = gerar_nome_saida(arquivo, "_converted", ".txt")
        saida.write_text(texto, encoding='utf-8')
        return saida
